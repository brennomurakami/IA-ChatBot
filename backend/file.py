from PyPDF2 import PdfReader
from langchain.text_splitter import TokenTextSplitter
from langchain_community.vectorstores import FAISS
import os
from dotenv import load_dotenv
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import docx
from langchain.docstore.document import Document
from io import BytesIO

load_dotenv()

embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")


def limpar_texto(texto):
    return " ".join(texto.split())


def process_document(files, file_extension):
    if file_extension == "pdf":
        print("PROCESSANDO PDF")
        reader = PdfReader(files)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                raw_text += text
    elif file_extension == "txt":
        print("PROCESSANDO TXT")
        if isinstance(files, BytesIO):
            files.seek(0)
            raw_text = files.read().decode("utf-8")
        else:
            with open(files, "r") as file:
                raw_text = file.read()
    elif file_extension == "docx":
        print("PROCESSANDO DOCX")
        doc = docx.Document(files)
        raw_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    elif file_extension == "string":
        return files
    else:
        raise ValueError("Unsupported file type")

    text_splitter = TokenTextSplitter(chunk_size=350, chunk_overlap=100)
    documents = text_splitter.split_text(raw_text)
    document_objs = [Document(page_content=doc) for doc in documents]
    print(document_objs)
    return document_objs

def process_document_return(files, file_extension):
    raw_text = " documento121: "
    if file_extension == "pdf":
        print("PROCESSANDO PDF")
        reader = PdfReader(files)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                raw_text += text
    elif file_extension == "txt":
        print("PROCESSANDO TXT")
        if isinstance(files, BytesIO):
            files.seek(0)
            raw_text = files.read().decode("utf-8")
        else:
            with open(files, "r") as file:
                raw_text = file.read()
    elif file_extension == "docx":
        print("PROCESSANDO DOCX")
        doc = docx.Document(files)
        raw_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    elif file_extension == "string":
        return files
    else:
        raise ValueError("Unsupported file type")

    text_splitter = TokenTextSplitter(chunk_size=350, chunk_overlap=100)
    documents = text_splitter.split_text(raw_text)
    document_objs = [Document(page_content=doc) for doc in documents]
    print(document_objs)
    return document_objs


def criar_indices_faiss(files, file_extension):
    documents = process_document(files, file_extension)
    vector_store = FAISS.from_documents(documents, embeddings)
    vector_store.save_local("faiss_index")
    return vector_store


def carregar_indices_faiss():
    vector_store = FAISS.load_local(
        "faiss_index", embeddings, allow_dangerous_deserialization=True
    )
    return vector_store


def adicionar_texto_ao_indice(vector_store, files, file_extension):
    documents = process_document(files, file_extension)
    # Ensure embeddings dimensionality matches the index
    document_embeddings = embeddings.embed_documents(
        [doc.page_content for doc in documents]
    )
    if document_embeddings and len(document_embeddings[0]) == vector_store.index.d:
        vector_store.add_documents(documents)
        vector_store.save_local("faiss_index")
    else:
        print(
            f"Dimension mismatch: embeddings dimension {len(document_embeddings[0])} vs index dimension {vector_store.index.d}"
        )


def verificar_e_atualizar_indice(files, file_extension):
    if os.path.exists("faiss_index"):
        vector_store = carregar_indices_faiss()
        adicionar_texto_ao_indice(vector_store, files, file_extension)
    else:
        vector_store = criar_indices_faiss(files, file_extension)
    return vector_store


def procurar_similaridade(consulta):
    print(f"Embedding consulta: {consulta}")
    # consulta = consulta.replace("VECTOR121:", "").strip()
    # query_embedding = embeddings.embed_query(consulta)
    vector_store = carregar_indices_faiss()
    resultados = vector_store.similarity_search(query=consulta, k=1000)
    textos_resultados = [limpar_texto(doc.page_content) for doc in resultados]
    textos_resultados = "".join(textos_resultados)
    print("\n\n\n", textos_resultados, "\n\n")
    return textos_resultados
